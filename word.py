#利用python查找word文档中的关键词，支持多个文档和多个关键词

# 导入所需库
import os,re
from docx import Document
global ur
 
ur = "D:/testPy" #全局变量读取输入路径
# key_word = '模糊查询' #全局变量读取关键字
 
 #获取文件路径下所有的word文件
def get_doc_path (path) :
    # file_list存储文件夹下所有文件
    file_list = os.listdir(path)
    # 正则匹配路径下所有.docx结尾的文件
    doc_list = [i for i in file_list if re.compile(r'\w+.docx').match(i)]
    #拼接ur进入doc_list，获得完整路径
    for i in range(len(doc_list)):
        doc_list[i] = ur + '/' + doc_list[i]
    return doc_list

# #获取路径下所有pdf文件
# def get_pdf_path (path) :
#     # file_list存储文件夹下所有文件
#     file_list = os.listdir(path)
#     # 正则匹配路径下所有.pdf结尾的文件
#     pdf_list = [i for i in file_list if re.compile(r'\w+.pdf').match(i)]
#     #拼接ur进入pdf_list，获得完整路径
#     for i in range(len(pdf_list)):
#         pdf_list[i] = ur + '/' + pdf_list[i]
#     return pdf_list


#查找文件中关键字所在行，参数为文件路径以及关键字
def find_word_text(path,word):
    fp = open(path, "r", encoding="utf-8")
    line_list = fp.readlines()
    for line in data:
        Code_line = line_list.index(line) + 1
        Now_code = line.strip("\n")
        print("path")
        print("----------------------------------------")
        for unsafe in [word]:
                flag = re.findall(unsafe, Now_code)
                if len(flag) != 0:
                    print("行号: {}    关键字内容行: {} " .\
                          format(Code_line,item))


# #查找关键字，函数值为文件路径以及关键字
# def find_text (path,word):
#     #获取word文件
#     document = Document(path)
#     #返回word文件所有段落集合
#     all_paragraphs = document.paragraphs
#     list1 = []
#     for paragraph in all_paragraphs:
#         str1 = paragraph.text
#         if str1.find(word) != -1 :
#             list1.append(str1)
#     if list1 == [] :
#         list1.append('未找到关键字')
#     return list1
#         # else:print('notfind')
 
 
# #获取路径下txt中的关键字并返回一个list
# def read_keyword (add) :
#     with open(add+'/key_word.txt',encoding='utf=8') as f :
#         key_word = f.readlines()
#         for i in range(len(key_word)) :
#             key_word[i] = key_word[i].strip()#去除换行符
#         return key_word
 
 
#获取文件夹路径

doc_path_list = get_doc_path(ur)
pdf_path_list = get_pdf_path(ur)
for item_doc in doc_path_list:
    find_word_text(item_doc,"你好")
for item_pdf in pdf_path_list:
    find_word_text(item_pdf,"你好")
# list_word = []
# document = Document()
# for x in doc_path :
#     document.add_heading(x,level=1)
#     for i in key_word :
#         document.add_paragraph(i , style='Intense Quote')
#         for num in range(len(find_text(x,i))):
#             no = num+1
#             str(no)
#             str1 = 'NO.' + str(no)
#             document.add_paragraph(str1)#添加标号
#             document.add_paragraph(find_text(x,i)[num])#写入正文
# document.save(ur + '/result.docx')
