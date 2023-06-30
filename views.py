from django.shortcuts import render
from django.http import JsonResponse
import json
from django.views.decorators.csrf import csrf_exempt

import tkinter as tk
from tkinter import filedialog

import os
import PyPDF2
from win32com import client as wc
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx import Document

# Create your views here.

#加载页面
def index(req):
    return render(req,'index.html')


index_file = "file_index.txt"
selected_results = []


# 建立文件索引并存储在文件index_file中
def build_file_index(directory):
    with open(index_file, "w") as file:
        for root, dirs, files in os.walk(directory):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                file_type = get_file_type(file_name)
                if file_type in ['pdf', 'word']:
                    file_info = f"{file_name}\t{file_path}\t{file_type}\n"
                    file.write(file_info)


# 获取文件类型
def get_file_type(file_name):  
    # 根据文件名后缀判断文件类型
    file_extension = os.path.splitext(file_name)[1].lower()
    if file_extension == '.pdf':
        return 'pdf'
    elif file_extension in ['.doc', '.docx']:
        return 'word'
    else:
        return 'other'


# 关键字查询
def search_keywords(keywords):
    matching_results = []
    with open(index_file, "r") as file:
        current_file_path = None
        line_number = 0
        for line in file:
            file_name, file_path, file_type = line.strip().split('\t')
            if file_type in ['pdf', 'word']:
                if current_file_path != file_path:
                    current_file_path = file_path
                    line_number = 0
                else:
                    line_number += 1
                matches = find_keyword_matches(keywords, file_name, file_path, line_number)
                matching_results.extend(matches)
    return matching_results


# 在文件中查找关键字匹配
def find_keyword_matches(keywords, file_name, file_path, line_number):
    matches = []
    # 如果是doc文件则转换成为docx文件
    # if file_path.endswith('.doc'):
    #     doc_to_docx(file_path)
    # pdf处理
    if file_path.endswith('.pdf'):
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_number = 0
            for page_number in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_number]
                text = page.extract_text()
                lines = text.split('\n')
                for line in lines:
                    line_number += 1
                    if any(keyword in line for keyword in keywords):
                        match = (file_path, line_number, line.strip())
                        matches.append(match)
    # word处理
    elif file_path.endswith('.docx'):
        try:
            matches = []
            doc = Document(file_path)

            # 处理标题
            if isinstance(doc.element.body, Paragraph) and doc.element.body.style.name == 'Title':
                title = doc.core_properties.title  # 获取文档标题
                line_number = 1  # 标题行号为1
                if any(keyword in title for keyword in keywords):
                    match = (file_path, line_number, title.strip())
                    matches.append(match)
            
            # 处理正文文本内容  
            for paragraph_number, paragraph in enumerate(doc.paragraphs, 1):
                line_number += 1  # 以段落号作为行号
                for run in paragraph.runs:  # 分割行
                    if any(keyword in run.text for keyword in keywords):
                        match = (file_path, line_number, run.text.strip())
                        matches.append(match)

            # 处理表格内内容，包含合并单元格表格的处理     
            for element in doc.element.body:
                if isinstance(element, CT_Tbl):
                    table_contents = extract_table_content(file_path)
                    for table_data in table_contents:
                        for row in table_data:
                            line_number += 1  # 表格行号作为行号
                            if any(keyword in cell_text for cell_text in row for keyword in keywords):
                                match = (file_path, line_number, '\t'.join(row))
                                matches.append(match)
        except Exception as e:
            print(f"Error occurred while processing '{file_name}': {e}")
    return matches


# doc文件转docx文件
def doc_to_docx(file_path):
    word = wc.Dispatch("Word.Application")  # 打开word应用程序
    # for file in files:
    doc = word.Documents.Open(file_path)  # 打开word文件
    doc.SaveAs("{}x".format(file_path), 12)  # 另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close()  # 关闭原来word文件
    word.Quit()
    # print("完成！")


# 处理docx文件中表格信息，合并单元格的单元格只输出一次
def extract_table_content(doc_path):
    doc = Document(doc_path)
    table_contents = []
    for table in doc.tables:
        table_data = []
        cells = table._cells
        rows = table.rows
        cols = table.columns
        for i, cell in enumerate(cells):
            if cell in cells[:i]:  # 如果该单元格不是在表中第一次出现则跳过
                continue
            for j in range(len(rows)):
                if i >= j * len(cols) and i < (j + 1) * len(cols):
                    row_index = j
                    # col_index = i % len(cols)
                    break
            cell_text = cell.text
            if is_merged_cell(cells, i):  # 判断是否为合并单元格
                cell_text = ''
            while len(table_data) <= row_index:
                table_data.append([])
            table_data[row_index].append(cell_text)

        table_contents.append(table_data)

    return table_contents


# 判断是否为合并单元格
def is_merged_cell(cells, index):
    cell = cells[index]
    for i in range(index - 1, -1, -1):  # 倒序查找
        if cell is cells[i]:  # 找到"相同"的单元格，如果没有"合并"单元格，则会倒序找到"自己"
            return True
    return False


# 用户选择结果(默认全部)
def select_result(result):
    selected_results.append(result)


# 用户保存结果到文件
def save_results_to_file():
    global keyfile
    keyfile={'res':''}
    if not selected_results:
        print("No result.")
        return
    
    with open("saved_results.txt", "w", encoding="utf-8") as file:  # 指定编码为UTF-8
        
        current_file_path = None
        res=""
        for file_path, line_number, content in selected_results:
            if current_file_path != file_path:
                current_file_path = file_path
                file.write(f"{file_path}\n{'-' * 30}\n")
                res+=f"*{file_path}@"
            file.write(f"行号{line_number}\t{content}\n")
            res+=f"行号{line_number}&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;\t{content}@"
        keyfile={'res':res}
    
    print("Saved as saved_results.txt")


# 主程序 需要前端提供 directory 与 keywords
def init(directory,keywords):
    # directory = "E:/其他/通知"  # 替换为所需路径
    # keywords = ["学院"]  # 替换为所需关键词
    build_file_index(directory)
    search_results = search_keywords(keywords)
    for result in search_results:
        select_result(result)
    save_results_to_file()

@csrf_exempt
def get_folder_path(request):
    if request.method == 'POST':
        #打开文件对话框
        root=tk.Tk()
        root.withdraw()
        root.wm_attributes("-topmost", 1)
        #获取文件夹路径
        folder_path=filedialog.askdirectory()

        data={'folder_path':folder_path}
        file_data=""
        #建立索引
        for root, dirs, files in os.walk(folder_path):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                file_type = get_file_type(file_name)
                if file_type in ['pdf', 'word']:
                    file_info = f"{file_path}*"
                    file_data+=file_info
        data.update({'files': file_data}) 
    return JsonResponse(data)



@csrf_exempt
def seek_keyword(request):
    if request.method=='POST':
        # global keywords 
        # global directory
        keywords = []  

        keywords.append(request.POST.get('keyword'))
        directory = request.POST.get('folder_path')
        selected_results.clear()
        # data = json.loads(request.body)
        # keywords= data.get('keyword')
        # directory = data.get('folder_path')
        init(directory,keywords)
    return  JsonResponse(keyfile)

@csrf_exempt
def download(request):
    res={'code':'0'}
    if request.method=='POST':
        text=request.POST.get('text')
        # 替换换行符
        text = text.replace('\n', '\n\n')
        text = text.replace('.pdf', f".pdf\n{'-' * 80}")
        text = text.replace('.docx', f".docx\n{'-' * 80}")
        text = text.replace('行号', '\n行号')

        # 保存文件名的变量
        filename = request.POST.get('key')+".txt"

        # 指定保存路径
        # path = "D:/Download/"
        #打开文件对话框
        root=tk.Tk()
        root.withdraw()
        root.wm_attributes("-topmost", 1)
        #获取文件夹路径
        path=filedialog.askdirectory()
        if path:
            # 用户选择了文件夹路径
            file_path = os.path.join(path + "/", filename)
        else:
             return JsonResponse(res)

        # 将文本写入文件
        with open(file_path, "w", encoding="utf-8") as file:
            file.write(text)
        res={'code':'1'}
    return JsonResponse(res)
